import os
import sys

# Add user site packages to sys.path
sys.path.append(r'C:\Users\ADMIN\AppData\Roaming\Python\Python311\site-packages')

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Table, TableStyle, PageBreak, HRFlowable
)
from reportlab.pdfgen import canvas

DOCX_OUTPUT_PATH = r"c:\Users\ADMIN\Documents\CareerLaunch2\KnowledgePaat_Safe_Project_File_Cleanup_Audit.docx"
PDF_OUTPUT_PATH = r"c:\Users\ADMIN\Documents\CareerLaunch2\KnowledgePaat_Safe_Project_File_Cleanup_Audit.pdf"

print("Starting Cleanup Audit Document Generation...")

# ==========================================
# 1. GENERATE DOCX DOCUMENT
# ==========================================
def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'))

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

doc = Document()

for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

# Title Cover
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(28)
title_p.paragraph_format.space_after = Pt(6)
title_run = title_p.add_run("KNOWLEDGEPAAT")
title_run.font.name = "Calibri"
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(30, 58, 138)

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_after = Pt(14)
sub_run = sub_p.add_run("Safe Project File Cleanup & Repository Health Audit")
sub_run.font.name = "Calibri"
sub_run.font.size = Pt(15)
sub_run.font.color.rgb = RGBColor(13, 148, 136)

meta_p = doc.add_paragraph()
meta_p.paragraph_format.space_after = Pt(20)
meta_run = meta_p.add_run("Status: Analysis & Recommendation Report (Zero Destructive Actions)\nDate: August 29, 2026 | Mode: Non-Destructive Review Gate\nRepository: KnowledgePaat (CareerLaunch2)")
meta_run.font.name = "Calibri"
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(71, 85, 105)

doc.add_page_break()

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4.5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(15, 23, 42)
    return p

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.12
    r_bold = p.add_run(bold_prefix + ": ")
    r_bold.font.name = "Calibri"
    r_bold.font.size = Pt(10)
    r_bold.font.bold = True
    r_bold.font.color.rgb = RGBColor(15, 23, 42)
    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(15, 23, 42)
    return p

# Section 1
add_h1("1. Executive Summary")
add_body("This document presents the comprehensive, non-destructive Safe Project File Cleanup Audit for the KnowledgePaat codebase. Every file across the repository was inspected against imports, references, package configurations, database migrations, and runtime dependencies.")
add_body("Key Findings: Total tracked files = 162. Untracked files/artifacts = 43. Zero production source files are broken or unused. Heavy load-test binaries (132.8 MB in load-tests/bin/) and identical duplicate .doc reports were identified for local exclusion and cleanup upon user approval.")

# Section 2
add_h1("2. Total Project File Overview")
overview_data = [
    ["Directory / Module", "Count / Size", "Classification", "Purpose & Verified Role"],
    ["frontend/app/ (64 Routes)", "64 Page/Layouts", "🟢 REQUIRED", "Next.js 16.3 App Router Pages & Layouts"],
    ["frontend/app/api/ (15 APIs)", "15 Handlers", "🟢 REQUIRED", "Active Server Route Handlers (100% Production Backend)"],
    ["frontend/components/", "10 Components", "🟢 REQUIRED", "Reusable UI Component Library"],
    ["frontend/lib/", "16 Utilities", "🟢 REQUIRED", "Core Modules: SWR Cache, Rate Limiter, Logger, AI"],
    ["frontend/hooks/ & context/", "6 Modules", "🟢 REQUIRED", "useClientQuery, FeatureFlags, ThemeContext"],
    ["frontend/scripts/", "23 Test Scripts", "🟡 LOCAL DEV", "Development verification & seed scripts"],
    ["load-tests/scenarios/ & config/", "5 Scripts", "🟢 REQUIRED", "k6 Load Testing Scenario Suite"],
    ["load-tests/bin/", "132.8 MB (Binaries)", "🟡 LOCAL ONLY", "k6 Windows Executable Binaries"],
    ["scripts/", "2 Generator Scripts", "🟡 LOCAL ONLY", "Docx and PDF Document Generator Utilities"],
    ["backend/ (FastAPI Prototype)", "1 Service", "🟠 LEGACY", "Inactive prototype (0 production traffic)"],
    ["Root SQL Migrations", "33 SQL Files", "🟢 MASTER / 🟠 STEPS", "Master Schema & Step Setup Scripts"],
    ["Root Audit Reports & Docs", "12 Documents", "🟡 LOCAL ONLY", "Word, PDF & Markdown Architecture Reports"]
]

t_over = doc.add_table(rows=len(overview_data), cols=4)
t_over.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row in enumerate(overview_data):
    for c_idx, val in enumerate(row):
        cell = t_over.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
        set_cell_margins(cell, top=50, bottom=50, left=70, right=70)
        if r_idx == 0:
            set_cell_background(cell, "1E3A8A")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            if r_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 3: Required Files
add_h1("3. Required Project Files (Status: KEEP)")
add_bullet("Frontend Core", "package.json, next.config.ts, tsconfig.json, globals.css, layout.tsx, page.tsx, error.tsx, global-error.tsx, robots.ts, sitemap.ts.")
add_bullet("All 64 Application Pages", "Student Portal (dashboard, jobs, notes, prep, mock-interview), Admin Portal, Public Website.")
add_bullet("15 Route Handlers", "All endpoints in frontend/app/api/* (AI, payments, progress, rate-limited utilities).")
add_bullet("Core Libraries", "lib/clientQueryCache.ts, lib/rateLimit.ts, lib/logger.ts, lib/subscription.ts, lib/careerProgress.ts, lib/ai/*.")
add_bullet("Database Schemas", "careerlaunch_master_setup.sql (Master Schema) and supabase_performance_indexes.sql (18 Performance Indexes).")
add_bullet("Load Testing Source", "load-tests/config/environment.js, load-tests/scenarios/*.js.")

# Section 4: Generated Files
add_h1("4. Generated Files (Status: KEEP LOCAL / IGNORE RECOMMENDED)")
add_bullet("frontend/.next/", "Next.js Turbopack build cache and serverless bundles (Regenerated via npm run build).")
add_bullet("frontend/node_modules/", "NPM dependencies (Regenerated via npm install).")
add_bullet("backend/venv/", "Python virtual environment for FastAPI prototype.")
add_bullet("frontend/*.tsbuildinfo", "TypeScript incremental build info cache.")

# Section 5: Local-Only Files
add_h1("5. Local-Only Files (Status: KEEP LOCAL / DO NOT PUSH TO GITHUB)")
add_bullet("Large Binaries", "load-tests/bin/k6.exe (66.4 MB) - required for running k6 locally but should not bloat Git repo.")
add_bullet("Audit Deliverables", "KnowledgePaat_Final_Implementation_Architecture_Security_Performance_Analysis.docx and .pdf.")
add_bullet("Dev Scripts", "scripts/generate_final_audit_documents.py, frontend/scripts/*.ts.")
add_bullet("Brand Assets", "brand_assets/ (Raw image mockups and screenshots).")

# Section 6 & 7: Potentially Unused & Legacy
add_h1("6. Potentially Unused Files (Status: MANUAL REVIEW REQUIRED)")
add_bullet("gradzenx_complete_supabase_master.sql (52.5 KB)", "Master SQL schema from previous project name before rebrand to KnowledgePaat.")
add_bullet("ADMIN_CONTROLLED_STUDENT_PORTAL_DOCUMENTATION.md (12.5 KB)", "Operational manual for admin features.")

add_h1("7. Legacy Files (Status: REVIEW / ARCHIVE RECOMMENDED)")
add_bullet("backend/ Directory", "FastAPI MVP prototype. 100% of production traffic uses Next.js Route Handlers in frontend/app/api/*.")
add_bullet("Individual Setup SQLs (31 files)", "Earlier step-by-step setup SQLs now unified into careerlaunch_master_setup.sql.")

# Section 8 & 9: Safe Cleanup Candidates & Duplicates
add_h1("8. Safe Cleanup Candidates (Status: APPROVAL REQUIRED BEFORE DELETION)")
add_bullet("KNOWLEDGEPAAT_ARCHITECTURE_REVIEW_REPORT.doc (18.4 KB)", "100% byte-identical duplicate of .docx file.")
add_bullet("KNOWLEDGEPAAT_CAPACITY_AND_PERFORMANCE_REPORT.doc (16.2 KB)", "100% byte-identical duplicate of .docx file.")
add_bullet("load-tests/bin/k6-v0.56.0-windows-amd64/", "Extracted directory containing duplicate k6.exe (66.4 MB). Primary resides in load-tests/bin/k6.exe.")

add_h1("9. Duplicate Files Analysis")
dup_data = [
    ["Original File", "Duplicate File", "Similarity", "Action Recommendation"],
    ["KNOWLEDGEPAAT_ARCHITECTURE_REVIEW_REPORT.docx", "KNOWLEDGEPAAT_ARCHITECTURE_REVIEW_REPORT.doc", "100% Byte-for-Byte", "Delete .doc duplicate upon approval"],
    ["KNOWLEDGEPAAT_CAPACITY_AND_PERFORMANCE_REPORT.docx", "KNOWLEDGEPAAT_CAPACITY_AND_PERFORMANCE_REPORT.doc", "100% Byte-for-Byte", "Delete .doc duplicate upon approval"],
    ["load-tests/bin/k6.exe (66.4 MB)", "load-tests/bin/k6-v0.56.0-windows-amd64/k6.exe", "100% Byte-for-Byte", "Delete extracted subfolder upon approval"]
]
t_dup = doc.add_table(rows=len(dup_data), cols=4)
t_dup.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row in enumerate(dup_data):
    for c_idx, val in enumerate(row):
        cell = t_dup.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
        set_cell_margins(cell, top=50, bottom=50, left=70, right=70)
        if r_idx == 0:
            set_cell_background(cell, "1E3A8A")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            if r_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Section 10 & 11: Dependencies & Large Files
add_h1("10. Unused Dependencies Analysis")
add_body("All 12 production packages in frontend/package.json (@supabase/ssr, @supabase/supabase-js, jsonwebtoken, lucide-react, mammoth, next, papaparse, pdf-parse, react, react-dom, xlsx, dotenv) are actively imported. Zero unused heavy packages detected.")

add_h1("11. Large Files Analysis")
add_bullet("load-tests/bin/k6.exe", "66.4 MB | k6 load engine executable. Keep locally; exclude via .gitignore.")
add_bullet("load-tests/bin/k6-v0.56.0-windows-amd64/", "66.4 MB | Redundant extracted archive folder. Delete upon approval.")
add_bullet("brand_assets/ChatGPT Image Aug 26...png", "934 KB | High-res logo mockup. Keep locally.")

# Section 12 & 13: Environment & Gitignore
add_h1("12. Environment and Secret Files")
add_bullet("frontend/.env", "Contains local Supabase publishable credentials. NOT TRACKED BY GIT.")
add_bullet("frontend/.env.example & backend/.env.example", "Safe configuration templates without real secrets. Properly tracked.")

add_h1("13. Proposed Root .gitignore Additions (FOR REVIEW ONLY — NOT APPLIED)")
gitignore_code = """# ==========================================
# Root .gitignore for KnowledgePaat
# ==========================================
# 1. Local Environment & Secrets
.env
.env*.local
.env.production
*.pem

# 2. Build Outputs & Next.js Caches
frontend/.next/
frontend/out/
frontend/build/
frontend/dist/
frontend/*.tsbuildinfo
frontend/next-env.d.ts

# 3. Dependencies & Virtual Environments
frontend/node_modules/
backend/venv/
backend/__pycache__/
*.pyc

# 4. Load Testing Binaries & Large Artifacts
load-tests/bin/
*.zip

# 5. Local Review Reports & Generated Binaries
*.docx
*.doc
*.pdf
scripts/test_doc_env.py

# 6. IDE & OS Metadata
.DS_Store
Thumbs.db
.vscode/
.idea/"""
add_body(gitignore_code)

# Section 14 & 15: Tracked Check & Recommendations
add_h1("14. Files Already Tracked But Recommended for Local-Only")
add_body("None. All large binaries (k6.exe) and review documents are currently untracked by Git, so no 'git rm --cached' command is necessary.")

add_h1("15. Recommended Cleanup Actions & Final Approval Gate")
add_bullet("Action 1 (Safe Deletion)", "Remove redundant exact duplicates: KNOWLEDGEPAAT_ARCHITECTURE_REVIEW_REPORT.doc, KNOWLEDGEPAAT_CAPACITY_AND_PERFORMANCE_REPORT.doc, and load-tests/bin/k6-v0.56.0-windows-amd64/.")
add_bullet("Action 2 (.gitignore Setup)", "Create the root .gitignore with rules from Section 13.")
add_bullet("Action 3 (Keep Untouched)", "Preserve all 64 frontend routes, 15 API Route Handlers, components, hooks, SWR caches, rate limiters, and SQL schemas.")

doc.save(DOCX_OUTPUT_PATH)
print(f"DOCX Generated: {DOCX_OUTPUT_PATH}")

# ==========================================
# 2. GENERATE PDF DOCUMENT
# ==========================================
print("Generating PDF...")

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        if self._pageNumber > 1:
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            self.drawString(54, 750, "KnowledgePaat — Safe Project File Cleanup Audit")
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(54, 742, 558, 742)
            self.line(54, 45, 558, 45)
            self.drawString(54, 32, "Confidential — File Cleanup Audit Report")
            self.drawRightString(558, 32, f"Page {self._pageNumber} of {page_count}")
        self.restoreState()

pdf = SimpleDocTemplate(
    PDF_OUTPUT_PATH,
    pagesize=letter,
    leftMargin=54,
    rightMargin=54,
    topMargin=54,
    bottomMargin=54
)

styles = getSampleStyleSheet()

pdf_title = ParagraphStyle('PDFTitle', fontName='Helvetica-Bold', fontSize=22, leading=26, textColor=colors.HexColor("#1E3A8A"), spaceAfter=5)
pdf_sub = ParagraphStyle('PDFSub', fontName='Helvetica', fontSize=12, leading=15, textColor=colors.HexColor("#0D9488"), spaceAfter=12)
pdf_meta = ParagraphStyle('PDFMeta', fontName='Helvetica', fontSize=9, leading=12.5, textColor=colors.HexColor("#475569"), spaceAfter=15)
pdf_h1 = ParagraphStyle('PDFH1', fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=colors.HexColor("#1E3A8A"), spaceBefore=12, spaceAfter=5, keepWithNext=True)
pdf_body = ParagraphStyle('PDFBody', fontName='Helvetica', fontSize=8.5, leading=11.5, textColor=colors.HexColor("#0F172A"), spaceAfter=4)
pdf_cell = ParagraphStyle('PDFCell', fontName='Helvetica', fontSize=7.5, leading=9.5, textColor=colors.HexColor("#0F172A"))
pdf_cell_hdr = ParagraphStyle('PDFCellHdr', fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white)

story = []

story.append(Paragraph("KNOWLEDGEPAAT", pdf_title))
story.append(Paragraph("Safe Project File Cleanup & Repository Health Audit", pdf_sub))
story.append(Paragraph("<b>Status:</b> Non-Destructive Analysis (0 Files Modified)<br/><b>Date:</b> August 29, 2026 | <b>Scope:</b> Full Repository Inventory & Health Review", pdf_meta))
story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1E3A8A"), spaceAfter=12))

story.append(Paragraph("1. Executive Summary", pdf_h1))
story.append(Paragraph("This audit inspects the entire KnowledgePaat repository. All 64 frontend routes, 15 API Route Handlers, components, SWR query cache, rate limiter, and SQL migrations are verified active. 132.8 MB of local load-test binaries (k6.exe) and identical duplicate .doc reports were identified for local exclusion via .gitignore upon user approval.", pdf_body))

story.append(Paragraph("2. Total Project File Overview Table", pdf_h1))
pdf_over_rows = [[Paragraph(c, pdf_cell_hdr) for c in overview_data[0]]]
for r in overview_data[1:]:
    pdf_over_rows.append([Paragraph(r[0], pdf_cell), Paragraph(r[1], pdf_cell), Paragraph(r[2], pdf_cell), Paragraph(r[3], pdf_cell)])

t_pdf_over = Table(pdf_over_rows, colWidths=[130, 75, 75, 224])
t_pdf_over.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
    ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
    ('TOPPADDING', (0, 0), (-1, -1), 3),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
]))
story.append(t_pdf_over)

story.append(Paragraph("3. Duplicate Files Analysis", pdf_h1))
pdf_dup_rows = [[Paragraph(c, pdf_cell_hdr) for c in dup_data[0]]]
for r in dup_data[1:]:
    pdf_dup_rows.append([Paragraph(r[0], pdf_cell), Paragraph(r[1], pdf_cell), Paragraph(r[2], pdf_cell), Paragraph(r[3], pdf_cell)])

t_pdf_dup = Table(pdf_dup_rows, colWidths=[150, 150, 75, 129])
t_pdf_dup.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
    ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
    ('TOPPADDING', (0, 0), (-1, -1), 3),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
]))
story.append(t_pdf_dup)

story.append(Paragraph("4. Recommended Actions & Approval Gate", pdf_h1))
story.append(Paragraph("• <b>Safe Deletions:</b> Delete exact duplicates (.doc reports and redundant extracted k6 archive folder).<br/>• <b>.gitignore:</b> Add proposed root rules to exclude local binaries (k6.exe) and local reports from Git.<br/>• <b>Keep Untouched:</b> All application source code, 64 pages, 15 APIs, hooks, components, and SQL migrations remain intact.", pdf_body))

pdf.build(story, canvasmaker=NumberedCanvas)
print(f"PDF Generated: {PDF_OUTPUT_PATH}")
